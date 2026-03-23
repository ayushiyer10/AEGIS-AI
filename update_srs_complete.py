#!/usr/bin/env python3
"""
Complete AEGIS SRS Document Update Script
Adds Sections 11 and 12 with all subsections and content
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys
import io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.5):
    """Set paragraph spacing and line spacing"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pSpacing = OxmlElement('w:spacing')
    pSpacing.set(qn('w:before'), str(int(before * 20)))
    pSpacing.set(qn('w:after'), str(int(after * 20)))
    pSpacing.set(qn('w:line'), str(int(line_spacing * 240)))
    pSpacing.set(qn('w:lineRule'), 'auto')
    pPr.append(pSpacing)

def add_heading2(doc, ref_para, text):
    """Add a Heading 2 paragraph"""
    h2 = ref_para.insert_paragraph_before(text, style='Heading 2')
    for run in h2.runs:
        run.font.bold = True
    set_paragraph_spacing(h2, before=6, after=6, line_spacing=1.5)
    return h2

def add_normal(doc, ref_para, text):
    """Add a normal paragraph"""
    p = ref_para.insert_paragraph_before(text, style='Normal')
    set_paragraph_spacing(p, before=0, after=6, line_spacing=1.5)
    return p

def add_bold_text(doc, ref_para, text):
    """Add bold text paragraph"""
    p = ref_para.insert_paragraph_before(text, style='Normal')
    for run in p.runs:
        run.font.bold = True
    set_paragraph_spacing(p, before=0, after=3, line_spacing=1.5)
    return p

def add_bullet(doc, ref_para, text):
    """Add a bulleted item"""
    bullet = ref_para.insert_paragraph_before(text, style='List Bullet')
    set_paragraph_spacing(bullet, before=0, after=3, line_spacing=1.5)
    return bullet

# Load the document
print("Loading document...")
doc = Document("C:/Users/Ayush Iyer/Desktop/Anti Cheat/docs/AEGIS_SRS_Document.docx")
print(f"Original document: {len(doc.paragraphs)} paragraphs")

# Find Appendix A
appendix_idx = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == 'Appendix A: System Diagrams' and 'Heading 1' in para.style.name:
        appendix_idx = i
        break

print(f"Appendix A at index: {appendix_idx}")

# Find TOC entry for Appendix A
toc_appendix_idx = None
for i, para in enumerate(doc.paragraphs):
    if i > 14 and i < 30:
        if 'Appendix A' in para.text and 'Heading' not in para.style.name:
            toc_appendix_idx = i
            break

print(f"TOC Appendix A at index: {toc_appendix_idx}")

# Update TOC
if toc_appendix_idx:
    toc_ref = doc.paragraphs[toc_appendix_idx]
    toc_ref.insert_paragraph_before('11. Testing & Quality Assurance', style='List Bullet')
    toc_ref.insert_paragraph_before('12. Future Scope & Enhancements', style='List Bullet')

# Reference paragraph for all insertions
appendix_para = doc.paragraphs[appendix_idx]

# Blank line before section 11
appendix_para.insert_paragraph_before('')

# ============ SECTION 11: TESTING & QUALITY ASSURANCE ============
print("Adding Section 11...")

# Section 11 heading
h1_11 = appendix_para.insert_paragraph_before('11. Testing & Quality Assurance', style='Heading 1')
for run in h1_11.runs:
    run.font.bold = True
set_paragraph_spacing(h1_11, before=12, after=12, line_spacing=1.5)

# 11.1
add_heading2(doc, appendix_para, '11.1 Test Framework & Methodology')
add_normal(doc, appendix_para, 'The AEGIS anti-cheat engine employs a comprehensive testing strategy built on industry-standard frameworks and methodologies. The primary testing framework is JUnit 5 (Jupiter), which provides modern annotation-driven testing capabilities integrated seamlessly with Spring Boot Test infrastructure. Mock testing is performed using MockMvc, allowing developers to test HTTP endpoints without deploying a full application server. Temporary directory management is handled through the @TempDir annotation, ensuring clean test isolation and automatic cleanup of test artifacts. Test data setup and teardown are performed through fixture methods and Spring lifecycle hooks, guaranteeing consistent test environments across all test executions.')

# 11.2
add_heading2(doc, appendix_para, '11.2 Test Cases')

# Test Case 1
add_bold_text(doc, appendix_para, 'TEST CASE 1: ingestUpdatesLivePayloadWithExpectedFields')
add_normal(doc, appendix_para, 'Purpose: Validate that POST /api/ingest endpoint correctly processes mouse movement data and returns properly structured response payloads.')
add_normal(doc, appendix_para, 'Setup: The test creates a test payload containing four consecutive mouse movement events with coordinates progressing from (100,100) to (160,120) to (260,150) to (380,180). Test data includes device ID "test-device-001", timestamps ranging from 1.00 to 1.15 seconds, and a keyboard event "w" recorded at timestamp 1.12 seconds to simulate player input during mouse movements.')
add_normal(doc, appendix_para, 'Method: POST /api/ingest endpoint receives the constructed payload.')
add_bold_text(doc, appendix_para, 'Expected Results:')

test1_results = [
    'HTTP 200 OK response confirming successful ingestion',
    'GET /api/live returns with deviceId field containing "test-device-001"',
    'Response contains mouse array with detected movement events',
    'Response contains keyboard array with detected key press events',
    'velocity field is numeric type and greater than zero',
    'velocityAvg field is numeric type',
    'velocitySmoothed field is numeric type',
    'cheatScore field is numeric type',
    'isCheater field is boolean type'
]

for result in test1_results:
    add_bullet(doc, appendix_para, result)

add_normal(doc, appendix_para, 'Assertions: Nine comprehensive JSON path assertions validate the complete response structure and data types. Status: PASSED')

# Test Case 2
add_bold_text(doc, appendix_para, 'TEST CASE 2: snapshotsEndpointReturnsPersistedLiveSnapshots')
add_normal(doc, appendix_para, 'Purpose: Verify that /api/live/snapshots endpoint correctly returns historical snapshot data persisted from previous ingestion events.')
add_normal(doc, appendix_para, 'Setup: The test submits two sequential payload ingests to establish historical data. PayloadA contains mouse movement from (0,0) to (50,40) with timestamps 1.00 to 1.05 seconds, persisted at timestamp 1.20. PayloadB contains mouse movement from (50,40) to (120,90) with timestamps 1.10 to 1.15 seconds, persisted at timestamp 1.30. This creates a historical record of two distinct movement patterns.')
add_normal(doc, appendix_para, 'Method: POST /api/ingest executed twice for both payloads, followed by GET /api/live/snapshots?limit=2 to retrieve the persisted historical data.')
add_bold_text(doc, appendix_para, 'Expected Results:')

test2_results = [
    'Both ingest requests return HTTP 200 OK response',
    'Snapshots endpoint returns array data structure',
    'Array contains at least two snapshot objects',
    'Each snapshot object has velocity field (numeric)',
    'Each snapshot object has mouse array with movement history',
    'Each snapshot object has keyboard array with input history'
]

for result in test2_results:
    add_bullet(doc, appendix_para, result)

add_normal(doc, appendix_para, 'Assertions: Seven JSON path assertions verify snapshot persistence and data integrity. Status: PASSED')

# Test Case 3
add_bold_text(doc, appendix_para, 'TEST CASE 3: registerHeartbeatAndDeviceListEndpointsWork')
add_normal(doc, appendix_para, 'Purpose: Validate the complete agent lifecycle including registration, heartbeat mechanisms, and device enumeration endpoints.')
add_normal(doc, appendix_para, 'Setup: Register an agent with comprehensive device information including device ID "device-laptop-01", hostname "Laptop-01", operating system Windows, version 1.0.0, and IP address 192.168.1.100. This simulates a real-world agent initialization.')
add_normal(doc, appendix_para, 'Method 1 (Registration): POST /api/agents/register with agent configuration payload.')
add_bold_text(doc, appendix_para, 'Expected Results (Registration):')

for result in ['HTTP 200 OK response', 'Response status field equals "registered"', 'Response deviceId field equals "device-laptop-01"']:
    add_bullet(doc, appendix_para, result)

add_normal(doc, appendix_para, 'Method 2 (Heartbeat): POST /api/agents/heartbeat using same agent payload to maintain connection.')
add_bold_text(doc, appendix_para, 'Expected Results (Heartbeat):')

for result in ['HTTP 200 OK response', 'Response status field equals "alive"']:
    add_bullet(doc, appendix_para, result)

add_normal(doc, appendix_para, 'Method 3 (Device List): GET /api/devices to enumerate all active devices.')
add_bold_text(doc, appendix_para, 'Expected Results (Device List):')

for result in ['HTTP 200 OK response', 'Returns array of device objects', 'First device in array has deviceId "device-laptop-01"', 'First device has online boolean field']:
    add_bullet(doc, appendix_para, result)

add_normal(doc, appendix_para, 'Assertions: Eight assertions across three distinct endpoints validate the complete agent lifecycle. Status: PASSED')

# 11.3
add_heading2(doc, appendix_para, '11.3 Test Execution & Coverage')
add_normal(doc, appendix_para, 'Test Framework: Spring Boot Test environment with MockMvc for HTTP endpoint testing without server deployment.')
add_normal(doc, appendix_para, 'Test Environment: Temporary directory management per test using @TempDir annotation ensures isolated test execution without side effects. LiveDataStore.resetForTests() method ensures clean state before each test execution, preventing data contamination between test cases.')
add_normal(doc, appendix_para, 'Test Data Isolation: Each test maintains complete isolation through snapshot path configuration via engine.live.snapshot.path system property. Test payloads are constructed with unique device identifiers to prevent cross-test interference.')
add_normal(doc, appendix_para, 'Total Test Cases: Three comprehensive integration tests covering critical workflows.')
add_bold_text(doc, appendix_para, 'Coverage Areas:')

for item in ['Data ingestion validation through /api/ingest endpoint', 'Snapshot persistence and historical data retrieval', 'Agent registration, heartbeat, and device enumeration']:
    add_bullet(doc, appendix_para, item)

add_normal(doc, appendix_para, 'Test Results: All tests GREEN/PASSING with 100% assertion success rate.')

# 11.4
add_heading2(doc, appendix_para, '11.4 Testing Approach')
add_normal(doc, appendix_para, 'Current Implementation: Unit testing of feature extraction logic within LiveDataStore.java (906 lines). Integration testing of API endpoints through MockMvc framework. Regression testing ensures Live API stability across varied payload configurations and mouse movement patterns. Manual testing of frontend UI/UX components with real backend connectivity.')
add_normal(doc, appendix_para, 'Proposed Enhancements: Load testing infrastructure to validate system behavior under 10,000+ concurrent device connections. End-to-end testing automation of complete agent-to-frontend workflows. Automated UI testing using Selenium or Cypress frameworks for React components. Chaos engineering testing to validate system resilience under failure conditions.')

# 11.5
add_heading2(doc, appendix_para, '11.5 Test Data Management')
add_normal(doc, appendix_para, 'Temporary Files: All temporary test files are automatically cleaned after test execution through Spring Test framework integration. Snapshot path configuration is managed through engine.live.snapshot.path system property, providing centralized test data storage control.')
add_normal(doc, appendix_para, 'Payload Formats: Test payloads are documented with actual JSON structures reflecting real agent output. Mock HTTP client is built into Spring Test environment, eliminating need for external test fixtures.')
add_normal(doc, appendix_para, 'Quality Metrics: Code coverage targets 80% for backend services. Critical path coverage (data ingestion, detection, and reporting) targets 95% coverage. All new features require corresponding test cases before code review approval.')

print(f"Section 11 complete, document now has {len(doc.paragraphs)} paragraphs")

# ============ SECTION 12: FUTURE SCOPE & ENHANCEMENTS ============
print("Adding Section 12...")

# Blank line
appendix_para.insert_paragraph_before('')

# Section 12 heading
h1_12 = appendix_para.insert_paragraph_before('12. Future Scope & Enhancements', style='Heading 1')
for run in h1_12.runs:
    run.font.bold = True
set_paragraph_spacing(h1_12, before=12, after=12, line_spacing=1.5)

# 12.A
add_heading2(doc, appendix_para, '12.A Machine Learning Integration')
add_bold_text(doc, appendix_para, '1. ML Model Enhancement')

for bullet in [
    'Current: Rule-based detection with LiveDataStore.java (906 lines)',
    'Future: Integrate isolation_forest.joblib (pre-trained model in ml-engine/)',
    'Expected Improvement: Detect novel cheating patterns not covered by rules',
    'Implementation: Add ML scoring layer alongside behavioral analysis',
    'Timeline: Phase 2 (2-3 months)',
    'Risk Mitigation: Ensemble approach (rules + ML combined scoring)'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, '2. Continuous Model Training')

for bullet in [
    'Current: Static pre-trained model',
    'Future: Implement feedback loop for model retraining',
    'Data Source: Confirmed cheater cases and manual reports',
    'Frequency: Weekly/monthly retraining cycles',
    'Expected Benefit: Improved detection accuracy for evolving cheating methods'
]:
    add_bullet(doc, appendix_para, bullet)

# 12.B
add_heading2(doc, appendix_para, '12.B Real-Time WebSocket Integration')
add_bold_text(doc, appendix_para, '1. Replace Polling with Streaming')

for bullet in [
    'Current: Frontend polls every 3.5 seconds (CPU/network overhead)',
    'Future: Implement WebSocket for real-time event streaming',
    'Expected Improvement: Sub-100ms latency vs 3.5s polling delay',
    'Components Affected: GamingDashboard.jsx, GlobalIntel.jsx, Backend',
    'Timeline: Phase 2 (1-2 months)'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, '2. Server-Sent Events (SSE) Alternative')

for bullet in [
    'Lighter than WebSocket for one-way streaming',
    'Better browser support and simpler server implementation',
    'Option for teams preferring HTTP-based messaging'
]:
    add_bullet(doc, appendix_para, bullet)

# 12.C
add_heading2(doc, appendix_para, '12.C Enhanced Detection Capabilities')
add_bold_text(doc, appendix_para, '1. Mouse-Keyboard Correlation Analysis')

for bullet in [
    'Current: Analyze mouse and keyboard independently',
    'Future: Cross-correlate mouse velocity spikes with keyboard input patterns',
    'Detect: Rapid input sequences (spam clicking) with impossible aim precision',
    'Expected Accuracy Gain: +3-5% detection precision'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, '2. Eye-Tracking Integration')

for bullet in [
    'Integrate eye-tracking hardware support',
    'Validate that eye gaze correlates with cursor movement',
    'Detect: Cursor movement without corresponding eye movement (bot indicator)',
    'Market Requirement: Professional esports tournaments'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, '3. Game-Specific Profiles')

for bullet in [
    'Current: Generic cheat detection regardless of game',
    'Future: Create profiles for CS2, Valorant, Apex, Fortnite, etc.',
    'Game-Specific Tuning: Different detection thresholds per game',
    'Timeline: Phase 3 (ongoing)'
]:
    add_bullet(doc, appendix_para, bullet)

# 12.D
add_heading2(doc, appendix_para, '12.D Security & Authentication')
add_bold_text(doc, appendix_para, '1. API Authentication & Authorization')

for bullet in [
    'Current: No API authentication (localhost only)',
    'Future: JWT-based token authentication for backend APIs',
    'Scope: Protect /api/ingest, /api/agents/* endpoints',
    'Implementation: Spring Security integration'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, '2. End-to-End Encryption')

for bullet in [
    'Current: Plain HTTP for agent-to-backend communication',
    'Future: TLS/SSL encryption for all network traffic',
    'Certificate Management: Self-signed for local, CA-signed for cloud'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, '3. Agent Code Signing')

for bullet in [
    'Current: No code verification for agent authenticity',
    'Future: Digital signature verification on agent updates',
    'Purpose: Prevent tampered or malicious agent versions'
]:
    add_bullet(doc, appendix_para, bullet)

# 12.E
add_heading2(doc, appendix_para, '12.E Multi-Platform Support')
add_bold_text(doc, appendix_para, '1. Linux Client Support')

for bullet in [
    'Current: Windows-only agent (input_capture.py uses Windows APIs)',
    'Future: Implement Linux input capture using X11/Wayland',
    'Timeline: Phase 2-3'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, '2. macOS Client Support')

for bullet in [
    'Current: No macOS implementation',
    'Future: Native Darwin SDK for keyboard/mouse interception',
    'Market Requirement: Mac gamers (though niche for competitive gaming)'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, '3. Mobile Gaming (iOS/Android)')

for bullet in [
    'Current: Not applicable (touch-based gaming has different vulnerabilities)',
    'Future: Dedicate detection engine for mobile touchscreen patterns',
    'Scope: Detect tap-sequence automation, gesture spoofing, auto-clickers'
]:
    add_bullet(doc, appendix_para, bullet)

# 12.F
add_heading2(doc, appendix_para, '12.F Advanced Analytics & Reporting')
add_bold_text(doc, appendix_para, '1. Forensic Analysis Dashboard')
add_normal(doc, appendix_para, 'Current: GlobalIntel provides overview statistics. Future: Deep-dive forensic tools for security analysts.')

for bullet in [
    'Heatmaps of suspicious mouse movements',
    'Frame-by-frame playback of user input',
    'Timeline visualization of detection events',
    'Peer comparison analytics'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, '2. Automated Ban Recommendations')

for bullet in [
    'Current: Scored as cheater but manual review needed',
    'Future: Confidence-based auto-ban system for high-certainty cases',
    'Confidence Levels: >95% auto-ban, 80-95% flagged for review, <80% monitor'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, '3. Appeal Workflow')

for bullet in [
    'Mechanism for users to appeal ban decisions',
    'Review queue with priority routing',
    'Evidence presentation system'
]:
    add_bullet(doc, appendix_para, bullet)

# 12.G
add_heading2(doc, appendix_para, '12.G Distributed & Cloud Deployment')
add_bold_text(doc, appendix_para, '1. Kubernetes Containerization')

for bullet in [
    'Current: Monolithic Spring Boot JAR',
    'Future: Docker containers + Kubernetes orchestration',
    'Benefit: Horizontal scaling, load balancing, auto-recovery'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, '2. Scalable Database Migration')

for bullet in [
    'Current: File-based JSONL storage (not suitable for scale)',
    'Future: PostgreSQL or MongoDB cluster for multi-node deployments',
    'Benefits: Query flexibility, indexing, replication, high availability'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, '3. Message Queue Integration')

for bullet in [
    'Implement Apache Kafka or RabbitMQ',
    'Decouple agent data ingestion from detection processing',
    'Timeline: Phase 3 (large-scale deployment)'
]:
    add_bullet(doc, appendix_para, bullet)

# 12.H
add_heading2(doc, appendix_para, '12.H Integration with Game Ecosystems')
add_bold_text(doc, appendix_para, '1. Steam Integration')

for bullet in [
    'Current: Standalone system',
    'Future: API integration with Steamworks for VAC (Valve Anti-Cheat) coordination',
    'Benefit: Unified ban enforcement across platforms'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, '2. Discord Rich Presence')

for bullet in [
    'Display cheat detection statistics in Discord',
    'Auto-notify admins of critical detections'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, '3. Twitch Integration')

for bullet in [
    'Flag suspicious plays in streamed gameplay',
    'Archive flagged moments for evidence review'
]:
    add_bullet(doc, appendix_para, bullet)

# 12.I
add_heading2(doc, appendix_para, '12.I Performance Optimization')
add_bold_text(doc, appendix_para, '1. GPU Acceleration')

for bullet in [
    'Current: CPU-based feature extraction',
    'Future: CUDA/OpenCL for feature computation on GPU',
    'Expected Improvement: 5-10x faster processing'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, '2. Edge Computing')

for bullet in [
    'Move lightweight detection to agent (local decision-making)',
    'Reduce server load and latency',
    'Timeline: Phase 3'
]:
    add_bullet(doc, appendix_para, bullet)

# 12.J
add_heading2(doc, appendix_para, '12.J Regulatory Compliance')
add_bold_text(doc, appendix_para, '1. GDPR Compliance')

for bullet in [
    'Current: Stores mouse/keyboard data indefinitely',
    'Future: Data retention policies (delete after X days)',
    'Consent management for data collection'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, '2. Audit Logging')

for bullet in [
    'Track all ban decisions with reasoning',
    'Maintain searchable audit trail for legal review',
    'Export evidence for dispute resolution'
]:
    add_bullet(doc, appendix_para, bullet)

# 12.K
add_heading2(doc, appendix_para, '12.K Advanced Visualization & UI')
add_bold_text(doc, appendix_para, '1. 3D Geographic Heatmap')

for bullet in [
    'Current: 2D bar chart of regions',
    'Future: 3D globe showing global cheat distribution',
    'Interactive zoom, drill-down by region/country'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, '2. Real-Time Alert System')

for bullet in [
    'Current: Manual dashboard refresh',
    'Future: Notification system (email, SMS, Slack, Discord)',
    'Alerts for: Instant-flag events, high-confidence detections, system failures'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, '3. Mobile Admin Application')

for bullet in [
    'iOS/Android apps for on-the-go monitoring',
    'Quick action buttons (temporary/permanent ban, review details)'
]:
    add_bullet(doc, appendix_para, bullet)

# 12.L
add_heading2(doc, appendix_para, '12.L Testing & QA Expansion')
add_bold_text(doc, appendix_para, '1. Automated UI Testing')

for bullet in [
    'Selenium/Cypress tests for React components',
    'Regression test suite for UI changes'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, '2. Load Testing')

for bullet in [
    'Simulate 10,000+ concurrent agents',
    'Stress test snapshot persistence at scale',
    'Benchmark: Sub-100ms API response times'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, '3. Chaos Engineering')

for bullet in [
    'Test system resilience under failure conditions',
    'Network outages, database unavailability, agent crashes'
]:
    add_bullet(doc, appendix_para, bullet)

# 12.M
add_heading2(doc, appendix_para, '12.M Documentation & Knowledge Base')
add_bold_text(doc, appendix_para, '1. Final Administrator Guide')

for bullet in [
    'Deployment procedures',
    'Configuration tuning',
    'Troubleshooting common issues'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, '2. Machine Learning Model Documentation')

for bullet in [
    'Model architecture (isolation forest details)',
    'Training data characteristics',
    'Performance metrics (precision, recall, F1-score)'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, '3. Integration Guides')

for bullet in [
    'Step-by-step integration for game developers',
    'API client libraries (JavaScript, Python, C#/.NET)'
]:
    add_bullet(doc, appendix_para, bullet)

# 12.N
add_heading2(doc, appendix_para, '12.N Timeline & Prioritization')
add_normal(doc, appendix_para, 'The AEGIS enhancement roadmap is organized into four strategic phases, each with defined objectives, resources, and success metrics.')

add_bold_text(doc, appendix_para, 'PHASE 2 (Next 3-6 months) - HIGH PRIORITY:')

for bullet in [
    'WebSocket real-time streaming',
    'ML model integration',
    'JWT authentication',
    'TLS/SSL encryption',
    'Linux agent support',
    'Automated UI testing'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, 'PHASE 3 (6-12 months) - MEDIUM PRIORITY:')

for bullet in [
    'Kubernetes deployment',
    'Database migration to PostgreSQL',
    'Advanced analytics dashboard',
    'Mobile app (admin)',
    'Game-specific profiles (5+ titles)',
    'Steam integration'
]:
    add_bullet(doc, appendix_para, bullet)

add_bold_text(doc, appendix_para, 'PHASE 4 (12+ months) - LONG-TERM:')

for bullet in [
    'GPU acceleration',
    'GDPR compliance suite',
    '3D visualization',
    'Edge computing',
    'Mobile gaming support'
]:
    add_bullet(doc, appendix_para, bullet)

add_normal(doc, appendix_para, 'These enhancements position AEGIS as the industry-leading anti-cheat solution with comprehensive feature coverage, enterprise-grade scalability, and compliance-ready architecture.')

# Save the document
print(f"\nDocument complete with {len(doc.paragraphs)} paragraphs")
print("Saving document...")

output_path = "C:/Users/Ayush Iyer/Desktop/Anti Cheat/docs/AEGIS_SRS_Document.docx"
doc.save(output_path)

print(f"SUCCESS! Document saved to: {output_path}")
print(f"Final size: {len(doc.paragraphs)} paragraphs")
