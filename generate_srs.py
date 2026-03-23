#!/usr/bin/env python3
"""Generate comprehensive AEGIS Anti-Cheat SRS document"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_line_spacing(paragraph, spacing_value):
    """Set 1.5 line spacing"""
    pPr = paragraph._element.get_or_add_pPr()
    pSpacing = OxmlElement('w:spacing')
    pSpacing.set(qn('w:line'), '360')  # 1.5 spacing
    pSpacing.set(qn('w:lineRule'), 'auto')
    pPr.append(pSpacing)

def add_heading(doc, text, level):
    """Add styled heading"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12 if level > 1 else 14)
        run.font.bold = True
    set_line_spacing(h, 360)
    return h

def add_para(doc, text, bold=False):
    """Add styled paragraph"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = bold
    set_line_spacing(p, 360)
    return p

# Create document
doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ===== TITLE PAGE =====
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run('AEGIS Anti-Cheat Engine\nSoftware Requirements Specification')
title_run.font.name = 'Times New Roman'
title_run.font.size = Pt(18)
title_run.font.bold = True
set_line_spacing(title_para, 360)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Version 1.0')
subtitle_run.font.name = 'Times New Roman'
subtitle_run.font.size = Pt(12)
set_line_spacing(subtitle, 360)

doc.add_paragraph()
doc.add_paragraph()

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run(f'Document Date: {datetime.now().strftime("%B %d, %Y")}')
date_run.font.name = 'Times New Roman'
date_run.font.size = Pt(11)
set_line_spacing(date_para, 360)

doc.add_page_break()

# ===== EXECUTIVE SUMMARY =====
add_heading(doc, 'Executive Summary', 1)

add_para(doc, 'The AEGIS Anti-Cheat Engine is a comprehensive, real-time detection system designed to identify and prevent various forms of cheating in online gaming environments. The system leverages advanced movement analysis, behavioral pattern recognition, and machine learning-inspired scoring algorithms to detect suspicious activity with minimal false positives. By continuously analyzing mouse movements, keyboard inputs, screen motion, and game context, AEGIS provides administrators with actionable intelligence and enforcement capabilities to maintain competitive integrity.')

add_para(doc, 'The architecture is built on a three-tier model consisting of a Python-based client agent that runs on player machines, a Spring Boot backend that processes and analyzes behavioral data, and a React-based web interface for monitoring and analytics. The system maintains a 15-second offline timeout for device presence tracking and employs sophisticated state machines to balance detection sensitivity with legitimate play patterns. Real-time data flows through multiple processing pipelines that extract movement features, compute threat signals, and generate comprehensive risk assessments.')

add_para(doc, 'This specification defines all functional requirements, system components, detection algorithms, API contracts, and deployment architectures necessary to implement and maintain the AEGIS platform. The document serves as the authoritative reference for developers, testers, system administrators, and security personnel involved in the anti-cheat initiative.')

doc.add_page_break()

# ===== TABLE OF CONTENTS =====
add_heading(doc, 'Table of Contents', 1)

toc_items = [
    '1. Introduction',
    '2. System Overview and Architecture',
    '3. Functional Requirements',
    '4. System Components',
    '5. Non-Functional Requirements',
    '6. Data Requirements and Flow',
    '7. Deployment Architecture',
    '8. Detection Algorithms',
    '9. External Interfaces and API Specification',
    '10. Constraints and Assumptions',
    'Appendix A: System Diagrams'
]

for item in toc_items:
    toc_para = doc.add_paragraph(item, style='List Bullet')
    for run in toc_para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    set_line_spacing(toc_para, 360)

doc.add_page_break()

# ===== INTRODUCTION =====
add_heading(doc, '1. Introduction', 1)

add_heading(doc, '1.1 Purpose and Scope', 2)
add_para(doc, 'The purpose of this Software Requirements Specification is to document all functional and non-functional requirements for the AEGIS Anti-Cheat Engine. The scope encompasses the detection system architecture, data processing pipelines, user interfaces, deployment mechanisms, and the comprehensive set of algorithms that enable real-time cheat detection across multiple gaming platforms.')

add_heading(doc, '1.2 Project Overview', 2)
add_para(doc, 'AEGIS operates as a distributed anti-cheat solution where client agents continuously monitor player input behavior and screen activity. These agents transmit telemetry data to a central backend engine that performs sophisticated analysis using multiple detection vectors. The system maintains historical snapshots in JSONL format for forensic analysis and trend visualization. Game administrators and security teams interact with the system through a web-based interface that provides real-time monitoring, statistical analysis, and enforcement options.')

add_heading(doc, '1.3 Intended Audience', 2)
add_para(doc, 'This specification is intended for software architects, backend developers, frontend developers, DevOps engineers, QA testers, and system administrators. It provides the necessary technical depth for implementation teams while remaining accessible to stakeholders reviewing project scope and capabilities.')

doc.add_page_break()

# ===== SYSTEM OVERVIEW =====
add_heading(doc, '2. System Overview and Architecture', 1)

add_heading(doc, '2.1 High-Level Architecture', 2)
add_para(doc, 'AEGIS follows a three-tier distributed architecture. The client tier (Python CaptureAgent) runs on player machines and captures input events. The application tier (Spring Boot backend) processes telemetry, performs feature extraction, executes detection algorithms, and maintains system state. The presentation tier (React web interface) provides real-time monitoring, analytics, and administrator controls.')

add_para(doc, 'Data flows bidirectionally. Agents send input events through HTTP POST /api/ingest every 500 milliseconds. The backend queries agent status through heartbeat mechanisms and maintains persistent JSONL snapshots. The web frontend polls APIs every 3.5 seconds to fetch live device states and historical data for dashboard visualization.')

add_heading(doc, '2.2 Component Interactions', 2)
add_para(doc, 'The CaptureAgent initializes with device identification (APPDATA/AEGIS/device.id), establishes server connectivity via automatic discovery and failover, and monitors foreground windows. When a game is detected, enhanced sampling begins. Mouse events (500-entry buffer) and keyboard events (200-entry buffer) are flushed every 500ms.')

add_para(doc, 'LiveDataStore maintains concurrent maps for device states and agent presence. Each ingest payload triggers extractFeatures (mouse movement analysis), computeCheatScore (weighted metric synthesis), and evaluateThreatSignal (pattern matching). Snapshots persist to JSONL files.')

add_para(doc, 'GlobalIntel frontend fetches /api/devices and /api/live/snapshots, transforms metrics into regional aggregations, detection method classifications, and weekly trends. Color-coding indicates risk levels: clear, suspect, high-risk, confirmed.')

add_heading(doc, '2.3 Technology Stack', 2)
add_para(doc, 'Backend: Spring Boot 2.7+, Jackson JSON, Java 11+. Frontend: React 18+, Vite, Recharts, Framer Motion, Axios. Agent: Python 3.8+, pynput, PIL, requests. Storage: JSONL snapshots in data/live-snapshots/. Communication: HTTP/REST with CORS.')

doc.add_page_break()

# ===== FUNCTIONAL REQUIREMENTS =====
add_heading(doc, '3. Functional Requirements', 1)

add_heading(doc, '3.1 User Roles', 2)
add_para(doc, 'Player: End-user running games with AEGIS agent installed. Agent operates transparently.')
add_para(doc, 'Administrator: Monitor threat intelligence, review histories, issue enforcement decisions.')
add_para(doc, 'Analyst: Investigate patterns, review snapshots, refine thresholds based on false positives.')
add_para(doc, 'DevOps: Deploy, configure, maintain backend infrastructure and agent distributions.')

add_heading(doc, '3.2 Key Use Cases', 2)

use_cases = [
    ('Agent Registration', 'Player downloads agent, generates device ID, registers via /api/agents/register, begins monitoring.'),
    ('Input Capture', 'During gameplay, agent captures mouse, keyboard, and screen motion at high frequency.'),
    ('Game Detection', 'Agent detects game windows via GetForegroundWindow with heuristics: size, keywords, exclusion lists.'),
    ('Feature Extraction', 'Backend computes velocity, acceleration, jerk, straightness, direction change, click rates from mouse sequences.'),
    ('Cheat Scoring', 'Algorithm synthesizes 11 weighted signals (burst ratio, acceleration, patterns, screen motion, etc) into 0-1 score.'),
    ('Pattern Detection', 'Detects snap-triggers, aim-lock, rage patterns, impossible speed (>11k px/s), input desync.'),
    ('Risk Assessment', 'State machine determines isCheater flag and risk level: clear → suspect → high-risk → confirmed.'),
    ('Snapshot Storage', 'Every update persists JSON to JSONL for forensic review and trend analysis.'),
    ('Dashboard Display', 'Administrators view global stats: regions, detections, bans, rates, heatmaps, method breakdown.'),
    ('Forensic Analysis', 'Analysts query device snapshots to review detection progression and evidence accumulation.'),
    ('Trend Analysis', '6-week aggregation shows weekly detection frequencies and emerging patterns.'),
    ('Heartbeat Tracking', 'Agents heartbeat every 5s; devices offline after 15s silence.')
]

for title, desc in use_cases:
    p = doc.add_paragraph(f'{title}: ', style='List Bullet')
    p_run = p.runs[0]
    p_run.font.bold = True
    p_run.font.name = 'Times New Roman'
    d_run = p.add_run(desc)
    d_run.font.name = 'Times New Roman'
    d_run.font.size = Pt(12)
    set_line_spacing(p, 360)

doc.add_page_break()

# ===== SYSTEM COMPONENTS =====
add_heading(doc, '4. System Components', 1)

add_heading(doc, '4.1 Backend Architecture', 2)
add_heading(doc, '4.1.1 Spring Boot Framework', 3)
add_para(doc, 'Spring Boot application (EngineApplication) runs on embedded Tomcat, typically port 8080. Provides REST endpoints for agent registration, heartbeat, data ingestion, and queries. CORS enabled globally for cross-origin requests.')

add_heading(doc, '4.1.2 REST API Endpoints', 3)

endpoints = [
    ('POST /api/ingest', 'Receives mouse events, keys, game state, screen motion. Triggers LiveDataStore.update() for processing.'),
    ('POST /api/agents/register', 'Initial agent registration with deviceId, hostname, os, appVersion, ip metadata.'),
    ('POST /api/agents/heartbeat', 'Periodic heartbeat every 5 seconds to maintain online status.'),
    ('GET /api/health', 'Health check returning "ENGINE_OK" for connectivity verification.'),
    ('GET /api/devices', 'Returns list of all devices with current metrics, sorted by lastSeen descending.'),
    ('GET /api/live', 'Query parameter ?deviceId optional. Returns live state for device.'),
    ('GET /api/live/snapshots', 'Query parameters: limit (default 50), deviceId optional. Returns JSONL snapshots.'),
    ('GET /api/devices/{deviceId}/live', 'Live state for specific device.'),
    ('GET /api/devices/{deviceId}/snapshots', 'Historical snapshots for specific device.')
]

for endpoint, desc in endpoints:
    p = doc.add_paragraph(f'{endpoint}: ', style='List Bullet')
    p_run = p.runs[0]
    p_run.font.bold = True
    p_run.font.name = 'Times New Roman'
    d_run = p.add_run(desc)
    d_run.font.name = 'Times New Roman'
    d_run.font.size = Pt(12)
    set_line_spacing(p, 360)

add_heading(doc, '4.1.3 Service Layer', 3)
add_para(doc, 'LiveDataStore (static utility): Maintains ConcurrentHashMap<String, DeviceState> for thread-safe state. Methods: update(payload), registerAgent(payload), heartbeat(payload), getLiveState(deviceId), getSnapshots(deviceId, limit), getDevices().')

add_para(doc, 'DeviceState (inner class): Encapsulates movement metrics, cheat scores, threat signals, and persistence logic. Implements synchronized update() for thread safety.')

add_heading(doc, '4.1.4 Data Persistence', 3)
add_para(doc, 'Snapshots stored as JSONL (newline-delimited JSON) in data/live-snapshots/{deviceId}.jsonl or custom path via system property. Each snapshot contains: timestamp, deviceId, mouse events, keyboard events, velocity metrics, cheat score, threat level, risk level, evidence score, game state, detection metadata. Concurrent writer pattern with SNAPSHOT_FILE_LOCK prevents file corruption.')

doc.add_page_break()

add_heading(doc, '4.2 Frontend Architecture', 2)
add_heading(doc, '4.2.1 React Components', 3)
add_para(doc, 'App.jsx: Root component with sidebar navigation, page switching between GamingDashboard and GlobalIntel. Framer Motion handles smooth transitions.')

add_para(doc, 'GlobalIntel.jsx: Primary analytics component. Fetches /api/devices and /api/live/snapshots every 3.5 seconds. Displays: global stats (regions, detections, bans, rate), regional bar chart, detection method pie chart, 6-week trend line chart, top 5 cheaters table.')

add_para(doc, 'GamingDashboard.jsx: Live device monitoring with real-time metrics display. Shows velocity, straightness, direction change, click frequency, threat levels, evidence progression.')

add_heading(doc, '4.2.2 State Management', 3)
add_para(doc, 'Component-level React hooks (useState). Polling via useEffect with 3.5 second interval. Ref-based animation targets for Anime.js entrance effects. No centralized Redux/Zustand store; polling ensures eventual consistency.')

add_heading(doc, '4.2.3 Data Visualization', 3)
add_para(doc, 'Recharts library provides: BarChart for regions, PieChart for detection methods, LineChart for 6-week trends. Interactive tooltips and active state highlighting. Custom renderActiveShape enhances pie slice visualization.')

add_heading(doc, '4.2.4 UI Features', 3)
add_para(doc, 'Anime.js drives staggered entrance animations for stat boxes, panels, offender items. Manual refresh button triggers data reload with visual feedback. Connection status indicator: "WORLDWIDE NETWORK" (online) or "ENGINE OFFLINE".')

doc.add_page_break()

add_heading(doc, '4.3 Python Agent', 2)
add_heading(doc, '4.3.1 Input Capture', 3)
add_para(doc, 'CaptureAgent class: Initializes with server URL resolution (cached in APPDATA/AEGIS/server.url). Server discovery: initial candidate → cached URL → localhost:8080 → LAN discovery. Fallback uses ThreadPoolExecutor with 64 workers probing network hosts.')

add_para(doc, 'Dual-mode input capture: pynput listeners for callback-based mouse/keyboard events, Windows polling loop (GetAsyncKeyState) fallback. Tracked keys: W, A, S, D, Space, Shift, Ctrl, Alt, R, 1-5. Polling: 16ms during game, 50ms idle.')

add_heading(doc, '4.3.2 Event Structures', 3)
add_para(doc, 'Mouse Event: {x, y, t, cx?, cy?, dx?, dy?, button?, pressed?}. Coordinates include absolute (x, y), center-relative (cx, cy), and delta (dx, dy).')

add_para(doc, 'Keyboard Event: {key, t}. Timestamp-tagged key identifier.')

add_para(doc, 'Ingest Payload: {deviceId, hostname, os, appVersion, ip, mouse, keys, game: {active, exe, title}, screen: {motion, sampleCount}, timestamp}')

add_heading(doc, '4.3.3 Game Detection', 3)
add_para(doc, 'Multi-factor heuristics: (1) Window size >=55% screen OR >=82% width AND >=72% height. (2) Game keywords: counter-strike, valorant, dota, fortnite, apex, pubg, steam, etc. (3) Exclusion list: vs, chrome, discord, etc. (4) Environment variable AC_GAME_EXE_NAMES for allowlist. When foreground window fails heuristics, game_active=false and buffering pauses.')

add_heading(doc, '4.3.4 Screen Motion', 3)
add_para(doc, 'PIL ImageGrab captures 96x96 pixels at game center every 220ms. Grayscale conversion, resize to 20x20. Motion computed as sum of absolute pixel differences / total pixels. 10-sample rolling average. Disabled if Pillow missing or non-Windows.')

doc.add_page_break()

add_heading(doc, '4.4 Detection Engine', 2)
add_heading(doc, '4.4.1 Feature Extraction', 3)
add_para(doc, 'extractFeatures() processes mouse sequences: pathDistance, duration, speedSum/speedSamples (mean speed), maxSpeed, maxAcceleration, maxJerk, straightness (netDisplacement/pathDistance), directionChangeRate, clickRate, snapClickRate.')

add_para(doc, 'Filters: dt < 3ms discarded, dist < 1.2 AND not clicked skipped, speed > 12k px/s rejected. Snap events: speed > 1400 AND dist > 6px. Snap-clicks: click within 120ms of snap.')

add_heading(doc, '4.4.2 Cheat Score Algorithm', 3)
add_para(doc, 'computeCheatScore synthesizes 11 signals: 16% burst ratio, 12% acceleration, 8% jerk, 12% pattern confidence, 8% straightness, 5% direction, 5% click, 5% trigger, 5% snap-trigger, 20% screen motion, 14% hyperspeed. Clamped [0,1]. Weighted sum × sample confidence. Baseline = max(120, max(velocityAverage, smoothedVelocity)). Pattern confidence reaches 1.0 when straightness > 0.988 AND directionChangeRate < 0.035 AND burstRatio > 3.4.')

add_heading(doc, '4.4.3 Threat Signals', 3)

signals = [
    ('Snap-Trigger', 'snapClickRate >= 0.45 AND speed > 1100 AND motion > 18. Score: 0.95. Instant flag: YES'),
    ('Aim-Lock', 'straightness > 0.985 AND directionChangeRate < 0.05 AND speed > 1300. Score: 0.88. Instant flag: YES'),
    ('Rage Pattern', 'cheatScore > 0.72 AND clickRate > 8.0 AND motion > 16. Score: 0.82. Instant flag: NO'),
    ('Jump Tracking', 'maxSpeed > 2500 AND motion > 18 AND samples >= 10. Score: 0.76. Instant flag: NO'),
    ('Impossible Speed', 'maxSpeed > 11000 AND samples >= 12. Score: 0.92. Instant flag: YES'),
    ('Input Desync', 'samples <= 2 AND motion > 22.0 (3+ consecutive). Score: 0.68. Instant flag: NO'),
    ('Elevated Motion', 'cheatScore >= 0.56 AND motion > 14. Score: 0.62. Instant flag: NO')
]

for sig_name, sig_desc in signals:
    p = doc.add_paragraph(f'{sig_name}: ', style='List Bullet')
    p_run = p.runs[0]
    p_run.font.bold = True
    p_run.font.name = 'Times New Roman'
    d_run = p.add_run(sig_desc)
    d_run.font.name = 'Times New Roman'
    d_run.font.size = Pt(12)
    set_line_spacing(p, 360)

add_heading(doc, '4.4.4 State Machine', 3)
add_para(doc, 'updateCheatState() manages detection via streak counters (suspiciousStreak, cleanStreak, modEvidenceStreak, screenInputDesyncStreak) and evidence accumulation. Risk levels: clear (evidence < 0.32), suspect (evidence >= 0.32), high-risk (isCheater OR threat >= 0.50), confirmed (instant-flag OR threat >= 0.80 OR evidence >= 0.88).')

add_para(doc, 'isCheater = true when: (1) instantFlag detected, (2) modEvidenceStreak >= 2 AND evidence >= 0.78 AND threat >= 0.78, (3) within 10s flag hold. Resets to false after: (1) 8+ no-sample ticks, OR (2) 4+ clean windows AND evidence <= 0.22.')

doc.add_page_break()

# ===== NON-FUNCTIONAL REQUIREMENTS =====
add_heading(doc, '5. Non-Functional Requirements', 1)

add_heading(doc, '5.1 Performance', 2)
add_para(doc, 'Agent: 16ms polling during game (62.5 Hz). 500ms batch transmission. Memory <50MB. Backend: <10ms per ingest, <50ms getDevices() for 100+ devices. Frontend: <200ms fetch+render cycle via 3.5s polls.')

add_heading(doc, '5.2 Scalability', 2)
add_para(doc, 'Stateless REST endpoints enable horizontal scaling. ConcurrentHashMap supports thousands of concurrent devices. JSONL append-only scales linearly. Frontend handles 10k+ device lists.')

add_heading(doc, '5.3 Reliability', 2)
add_para(doc, 'Agent survives temporary network outages via local buffering and cached server URLs. 15-second heartbeat timeout. Backend snapshots survive restarts via disk JSONL. No single point of failure in agent tier.')

add_heading(doc, '5.4 Security', 2)
add_para(doc, 'CORS enabled for web interface. No authentication in current version. Device ID as implicit identifier. Assume trusted network. Snapshot files should be OS-protected.')

add_heading(doc, '5.5 Resource Utilization', 2)
add_para(doc, 'Agent CPU: <5% during gameplay. Memory: <100MB with Python runtime. Network: 4-8 KB/s per agent. Backend CPU/disk proportional to device count.')

add_heading(doc, '5.6 Compatibility', 2)
add_para(doc, 'Agent: Windows 7+ (uses GetForegroundWindow, GetAsyncKeyState). Cross-platform pynput provides fallback. Backend: Java 11+. Frontend: Modern browsers (Chrome, Firefox, Safari, Edge).')

doc.add_page_break()

# ===== DATA REQUIREMENTS =====
add_heading(doc, '6. Data Requirements and Flow', 1)

add_heading(doc, '6.1 Data Dictionary', 2)
add_para(doc, 'Snapshot Object fields: timestamp (ms), deviceId, mouse events, velocity, velocityAvg, velocitySmoothed, keyboard events, cheatScore (0-1), suspiciousStreak (0-12), isCheater (bool), sampleCount, straightness (0-1), directionChangeRate, maxSpeed, clickRate, snapClickRate, screenMotion, threatLevel (0-1), threatReason (snap-trigger, aim-lock, rage-pattern, impossible-speed, input-desync, elevated-motion, none), riskLevel (clear, suspect, high-risk, confirmed), evidenceScore (0-1), modEvidenceStreak (0-8), gameActive (bool), gameExe, detectionCount, firstDetectedAt (ms), lastDetectedAt (ms), online (bool).')

add_heading(doc, '6.2 Data Flow', 2)
add_para(doc, 'Flow 1: Game active → Input capture → 500ms batch → POST /api/ingest → LiveDataStore.update() → Feature extraction → Cheat scoring → Snapshot persistence → Response.')

add_para(doc, 'Flow 2: Frontend 3.5s timer → GET /api/devices → Backend aggregates all DeviceStates → Returns JSON array → Frontend transforms to regions/methods → React renders → Charts animate.')

add_para(doc, 'Flow 3: Admin queries device → GET /api/devices/{id}/snapshots → Backend reads JSONL (last N lines) → JSON-deserialize → Return to frontend → Timeline visualization.')

add_heading(doc, '6.3 File Structures', 2)
add_para(doc, 'Agent: APPDATA/AEGIS/device.id (16-char hex), APPDATA/AEGIS/server.url (cached URL). Backend: data/live-snapshots/{deviceId}.jsonl (newline-delimited JSON, append-only).')

doc.add_page_break()

# ===== DEPLOYMENT =====
add_heading(doc, '7. Deployment Architecture', 1)

add_heading(doc, '7.1 Deployment Modes', 2)

modes = [
    ('Web Development', 'Frontend: npm run dev @ :5173. Backend: java -jar @ :8080. Vite proxies /api/* to backend. Hot reload enabled.'),
    ('Web Production', 'Frontend: npm run build → dist/. Copy to engine/src/main/resources/static/. Backend serves static + APIs. Single JAR deployment.'),
    ('Desktop GUI', 'Backend starts @ :8080. Static files in engine/target/classes/static/. Launcher opens localhost:8080 in browser. Appears as standalone app.'),
    ('Standalone Agent', 'python input_capture.py --server-url http://backend:8080. Discovers backend via LAN if needed. No frontend required.')
]

for mode, desc in modes:
    p = doc.add_paragraph(f'{mode}: ', style='List Bullet')
    p_run = p.runs[0]
    p_run.font.bold = True
    p_run.font.name = 'Times New Roman'
    d_run = p.add_run(desc)
    d_run.font.name = 'Times New Roman'
    d_run.font.size = Pt(12)
    set_line_spacing(p, 360)

add_heading(doc, '7.2 Build Processes', 2)
add_para(doc, 'Frontend: npm install → npm run build (output dist/, ~5 seconds, ~500 KB gzipped). Backend: mvn clean package (compile, test, bundle Spring Boot JAR, ~50-100 MB). Agent: pip install requirements.txt (pynput, requests, pillow). PyInstaller compiles to EXE (~50 MB disk, ~150 MB virtual).')

add_heading(doc, '7.3 Distribution Artifacts', 2)
add_para(doc, 'Backend JAR: engine-1.0.0.jar containing Spring Boot, services, controllers. Executable: java -jar engine-1.0.0.jar. Startup: ~10 seconds.')

add_para(doc, 'Frontend Bundle: dist/index.html, dist/assets/index-{hash}.js, dist/assets/style-{hash}.css. Served as static content.')

add_para(doc, 'Agent: input_capture.py or input_capture.exe. Launch with environment variables or CLI parameters.')

doc.add_page_break()

# ===== DETECTION ALGORITHMS =====
add_heading(doc, '8. Detection Algorithms', 1)

add_heading(doc, '8.1 Movement Analysis', 2)
add_para(doc, 'Analysis begins after 8+ calibration windows establish baseline velocity. Algorithm extracts velocity from pixel displacement and elapsed time. Burst events (>3x baseline) trigger secondary scoring. Direction consistency analyzed via angle derivatives. Legitimate players show natural curvature; aimbots show near-zero direction changes.')

add_para(doc, 'Example: Player moves 500px in 1s, baseline = 120 px/s. New burst = 1000 px/s → ratio = 8.33 → contributes 16% × clamp((8.33-2.2)/2.0) = 16% to score.')

add_heading(doc, '8.2 Click Patterns', 2)
add_para(doc, 'Click-rate normalized against 18 clicks/sec threshold. Snap-trigger flagged when click occurs within 120ms of rapid relocation (>1400 px/s). Snap-click ratio >0.35 indicates trigger assist. Normal: 2-6 clicks/sec. Trigger-assisted: 10-25 clicks/sec at snaps.')

add_para(doc, 'Example: 5 snap events in 10s, 4 clicks within 120ms of snaps → ratio = 4/5 = 0.80 → triggers condition: snapClickRate >= 0.45 AND speed > 1100 → instant flag, score 0.95.')

add_heading(doc, '8.3 Velocity Smoothing', 2)
add_para(doc, 'EMA with alpha=0.25 provides responsiveness and jitter dampening. Jerk (acceleration change rate) captures inhuman patterns. Legitimate: <100 px/s^3. Aimbots: >500 px/s^3.')

add_heading(doc, '8.4 Risk Scoring', 2)
add_para(doc, 'Cheat score: 0-1 synthesis of 11 weighted signals. Threat level: 78% previous + 42% cheatScore + 50% signalScore (exponential moving average). Evidence score: instant-flag +0.45; normal +0.62*cheatScore + 0.38*signal. Decays 4% active, 9% offline.')

add_heading(doc, '8.5 Decision Logic', 2)
add_para(doc, 'isCheater = true if: (1) instantFlag pattern, (2) modEvidenceStreak >= 2 AND evidence >= 0.78 AND streak >= 3 AND threat >= 0.78, (3) within flag hold. Resets if: (1) no samples 8+ ticks, OR (2) cleanStreak >= 4 AND evidence <= 0.22.')

add_para(doc, 'Real scenario: Macro moves 15k pixels instantly while clicking. Detection: maxSpeed > 11k ✓, samples >= 12 ✓ → impossibleSpeedPattern ✓ → instantFlag ✓ → isCheater true ✓ → riskLevel confirmed ✓ → enforcement triggered.')

doc.add_page_break()

# ===== EXTERNAL INTERFACES =====
add_heading(doc, '9. External Interfaces and API Specification', 1)

add_heading(doc, '9.1 REST API Detail', 2)
add_heading(doc, '9.1.1 Ingest Endpoint', 3)
add_para(doc, 'POST /api/ingest. Content-Type: application/json. Request: {deviceId, hostname, os, appVersion, ip, mouse: [{x, y, t, cx?, cy?, dx?, dy?, button?, pressed?}], keys: [{key, t}], game: {active, exe, title}, screen: {motion, sampleCount}, timestamp}. Response: {status: "received"}, HTTP 200. Triggers LiveDataStore.update().')

add_heading(doc, '9.1.2 Device List', 3)
add_para(doc, 'GET /api/devices. No query params. Response: [{deviceId, hostname, os, appVersion, ip, lastSeen, online, cheatScore, isCheater, ...}], HTTP 200. Sorted descending by lastSeen.')

add_heading(doc, '9.1.3 Live State', 3)
add_para(doc, 'GET /api/live. Query: ?deviceId optional. Response: {timestamp, deviceId, mouse, velocity, cheatScore, isCheater, ...}, HTTP 200.')

add_heading(doc, '9.1.4 Snapshots', 3)
add_para(doc, 'GET /api/live/snapshots. Query: limit (default 50, max 1000), deviceId optional. Response: [{snapshot}, ...] oldest-to-newest, HTTP 200. Reads last N lines from JSONL.')

add_heading(doc, '9.2 Error Handling', 2)
add_para(doc, 'Current: No explicit error codes; malformed requests silently ignored. Future: 400 Bad Request, 404 Not Found, 500 Internal Server Error with messages.')

add_heading(doc, '9.3 Data Formats', 2)
add_para(doc, 'JSON: Standard serialization via Jackson (backend) and native JS (frontend). JSONL: Newline-delimited JSON for snapshot persistence, enabling streaming parsers.')

doc.add_page_break()

# ===== CONSTRAINTS =====
add_heading(doc, '10. Constraints and Assumptions', 1)

add_heading(doc, '10.1 Technical Constraints', 2)
add_para(doc, 'Agent limited to 500 mouse, 200 key events in buffers. JSONL uncompressed; high-volume (10k devices × 6 weeks) consumes significant disk. No database; horizontal scaling requires load balancing. HTTP (not HTTPS); assumes trusted LAN. No request signing.')

add_heading(doc, '10.2 Environmental Constraints', 2)
add_para(doc, 'Agent: Windows 7+ required for full support (GetForegroundWindow, APIs). Linux/macOS degraded (no window detection, no screen capture). Python 3.8+ required. Backend: Java 11+. Port 8080 must be available. Filesystem must support concurrent appends.')

add_para(doc, 'Frontend: Modern browser (Chrome, Firefox, Safari, Edge). ES6+ support required. Static file serving via web server or embedded server.')

add_heading(doc, '10.3 Assumptions', 2)
add_para(doc, 'Players install agent cooperatively. Legitimate patterns statistically distinguishable from cheating. Administrators trusted; no audit log. Network bandwidth sufficient for 4-8 KB/s per agent. Detection thresholds may need tuning for new games/exploits.')

add_heading(doc, '10.4 Limitations', 2)
add_para(doc, 'Cannot detect cheats without client telemetry. Cannot prevent closed-source modifications. No cryptographic verification of input. Detection latency: 6s warmup + 10-20s for streak accumulation. False positives possible for accessibility tools. No ML adaptation; static thresholds. Dashboard 3.5s latency.')

doc.add_page_break()

# ===== APPENDIX: DIAGRAMS =====
add_heading(doc, 'Appendix A: System Diagrams', 1)

add_heading(doc, 'A.1 Architecture Diagram', 2)
arch = '''┌─────────────────────────────────────────┐
│         AEGIS THREE-TIER ARCHITECTURE    │
├──────────────┬──────────────┬────────────┤
│   CLIENT     │ APPLICATION  │PRESENTATION│
│              │              │            │
│ Python Agent │ Spring Boot  │ React Web  │
│ CaptureAgent │ Backend      │ GlobalIntel│
│ • Input Cap. │ • Controllers│ • Dashboard│
│ • Game Detect│ • Services   │ • 3.5s Poll│
│ • Event Buff │ • LiveDataStr│ • Charts   │
│ • Device ID  │ • File I/O   │            │
│              │              │            │
└────┬─────────┴────┬─────────┴─────┬──────┘
     │ POST /ingest │ GET /api      │
     │ 500ms batch  │ polls         │
     │ heartbeat 5s │ JSON response │
     │              │               │
     └──────────────┴───────────────┘
                    │
            ┌───────▼────────┐
            │ data/live-snap/│
            │ {id}.jsonl     │
            └────────────────┘
'''
add_para(doc, arch)

add_heading(doc, 'A.2 Data Processing Pipeline', 2)
pipeline = '''Input Events → Buffering → Payload → POST /api/ingest
                                               ↓
                                    Feature Extraction
                                    • Velocity, Accel
                                    • Straightness
                                    • Click Rates
                                               ↓
                                    Cheat Score (0-1)
                                    • 11 Signals
                                    • Weighted Synth
                                               ↓
                                    Threat Evaluation
                                    • Snap-Trigger?
                                    • Aim-Lock?
                                    • Impossible Speed?
                                               ↓
                                    State Machine
                                    • Streaks
                                    • Evidence
                                    • Risk Level
                                               ↓
                                    Snapshot JSONL
                                    └─ Forensics'''
add_para(doc, pipeline)

add_heading(doc, 'A.3 Risk Level State Diagram', 2)
states = '''Clear ← No Signals
         ↓ Suspicious Sample
      Suspect ← Evidence building
         ↓ Streak sustained
    High-Risk ← Threat accumulates
         ↓ Evidence >=0.78 OR Instant-Flag
    Confirmed ← isCheater = true
         ↓ Clean windows >= 4
      Clear ← Reset'''
add_para(doc, states)

doc.add_page_break()

# ===== APPENDIX B: TECHNOLOGY STACK =====
add_heading(doc, 'Appendix B: Technology Stack Summary', 1)

add_heading(doc, 'B.1 Technology Summary', 2)
add_para(doc, 'Backend Framework: Spring Boot 2.7+. Language: Java 11+. JSON: Jackson. Frontend Framework: React 18+. Build: Vite. Visualization: Recharts. Animation: Anime.js, Framer Motion. HTTP: Axios. Icons: Lucide React. Agent Language: Python 3.8+. Input: pynput. Capture: PIL/Pillow. HTTP: requests. Storage: JSONL (filesystem). Communication: HTTP/REST.')

doc.add_page_break()

# ===== CONCLUSION =====
add_heading(doc, 'Document Conclusion', 1)

add_para(doc, 'This Software Requirements Specification provides the authoritative foundation for AEGIS Anti-Cheat Engine implementation. The three-tier architecture, multi-vector detection approach, and sophisticated state machine represent a production-ready anti-cheat system.')

add_para(doc, 'The system successfully balances detection sensitivity against false positive reduction through evidence accumulation, modular streak tracking, and pattern-specific instant flags. Forensic capabilities via JSONL snapshots enable post-incident analysis and continuous refinement.')

add_para(doc, 'Future enhancements: encrypted HTTPS, database backend, ML model integration, cross-game pattern sharing, mobile support, user authentication.')

add_para(doc, 'This document is subject to revision as the project evolves.')

# Save
output_path = 'c:/Users/Ayush Iyer/Desktop/Anti Cheat/docs/AEGIS_SRS_Document.docx'
doc.save(output_path)

print(f'[OK] Document created: {output_path}')
print(f'[OK] Formatting: Times New Roman 12pt, 1.5 spacing')
print(f'[OK] Sections: Introduction through Constraints')
print(f'[OK] Diagrams: 3 ASCII diagrams in Appendix A')
print(f'[OK] Technology Stack: Appendix B')
print(f'[OK] Ready for submission')
